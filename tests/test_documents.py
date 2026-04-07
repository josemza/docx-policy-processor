def test_upload_docx_creates_initial_operation(client, auth_headers, sample_docx_bytes):
    products_response = client.get("/api/v1/products", headers=auth_headers)
    product_id = products_response.json()["data"][0]["id"]

    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        data={"product_id": product_id, "policy_number": "POL-001"},
        files={
            "file": (
                "prepoliza.docx",
                sample_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is True
    assert payload["data"]["status"] == "RECEIVED"
    assert payload["data"]["policy_number"] == "POL-001"
    assert payload["data"]["original_filename"] == "prepoliza.docx"
    assert payload["data"]["original_path"].startswith("documents/originals/")
    assert payload["data"]["output_path"].startswith("documents/outputs/")



def test_upload_rejects_non_docx(client, auth_headers):
    products_response = client.get("/api/v1/products", headers=auth_headers)
    product_id = products_response.json()["data"][0]["id"]

    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        data={"product_id": product_id, "policy_number": "POL-002"},
        files={"file": ("archivo.pdf", b"%PDF-1.4", "application/pdf")},
    )

    assert response.status_code == 400
    assert response.json()["error"]["code"] == "invalid_file_extension"



def test_process_docx_updates_header_title_and_preserves_table_and_emphasis(
    client,
    auth_headers,
    sample_docx_bytes,
    settings_obj,
):
    from docx import Document

    products_response = client.get("/api/v1/products", headers=auth_headers)
    product_id = products_response.json()["data"][0]["id"]

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        data={"product_id": product_id, "policy_number": "POL-003"},
        files={
            "file": (
                "prepoliza.docx",
                sample_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    operation = upload_response.json()["data"]

    process_response = client.post(
        f"/api/v1/documents/{operation['id']}/process",
        headers=auth_headers,
    )

    assert process_response.status_code == 200
    payload = process_response.json()
    assert payload["data"]["status"] == "COMPLETED"
    assert payload["data"]["error_message"] is None

    output_path = settings_obj.storage_root / payload["data"]["output_path"]
    assert output_path.exists()

    formatted = Document(output_path)
    assert formatted.sections[0].header.paragraphs[0].text == "Header POL-003"
    assert formatted.paragraphs[0].text == "TITULO VIDA TEST"
    body_run = formatted.paragraphs[1].runs[0]
    assert body_run.bold is True
    assert body_run.underline is True
    assert formatted.tables[0].cell(0, 0).text == "Celda 1"
    assert formatted.tables[0].cell(0, 1).text == "Celda 2"
    assert len(formatted.tables) == 1
