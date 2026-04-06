from collections.abc import Generator
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient


@pytest.fixture(autouse=True)
def test_environment(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> Generator[None, None, None]:
    db_path = tmp_path / "test_auth.db"
    storage_root = tmp_path / "storage"
    monkeypatch.setenv("ENVIRONMENT", "testing")
    monkeypatch.setenv("DATABASE_BACKEND", "sqlite")
    monkeypatch.setenv("DATABASE_SQLITE_PATH", str(db_path))
    monkeypatch.setenv("AUTH_BOOTSTRAP_ADMIN", "false")
    monkeypatch.setenv("BOOTSTRAP_PRODUCT_CATALOG", "false")
    monkeypatch.setenv("STORAGE_ROOT_PATH", str(storage_root))
    monkeypatch.setenv("DOCUMENTS_MAX_UPLOAD_BYTES", "1048576")

    from app.core.config import get_settings

    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


@pytest.fixture()
def client() -> Generator[TestClient, None, None]:
    import json

    from app.core.security import hash_password
    from app.infrastructure.db.base import Base
    from app.infrastructure.db.models.auth import UserModel
    from app.infrastructure.db.models.products import FormatRuleModel, ProductModel
    from app.infrastructure.db.session import get_session_local
    from app.main import create_app

    session_local, engine = get_session_local()
    Base.metadata.drop_all(bind=engine)
    Base.metadata.create_all(bind=engine)
    db = session_local()
    try:
        user = UserModel(
            username="admin",
            password_hash=hash_password("admin12345"),
            full_name="Admin Test",
            role="admin",
            is_active=True,
        )
        db.add(user)
        db.flush()
        product = ProductModel(
            code="VIDA_TEST",
            name="Vida Test",
            title_template="Titulo Vida Test",
            header_template="Header {{ numero_poliza }}",
            active=True,
        )
        db.add(product)
        db.flush()
        db.add(
            FormatRuleModel(
                product_id=product.id,
                version=1,
                configuration_json=json.dumps(
                    {
                        "page_setup": {
                            "paper_size": "A4",
                            "margin_top_cm": 2.5,
                            "margin_bottom_cm": 2.5,
                            "margin_left_cm": 2.0,
                            "margin_right_cm": 2.0,
                        },
                        "font_defaults": {"family": "Arial", "size_pt": 10},
                        "paragraph_defaults": {"line_spacing": 1.15, "alignment": "justify"},
                        "title_rules": {"alignment": "center", "bold": True, "case": "upper"},
                    }
                ),
                active=True,
            )
        )
        db.commit()
    finally:
        db.close()

    app = create_app()
    with TestClient(app) as test_client:
        yield test_client

    Base.metadata.drop_all(bind=engine)


@pytest.fixture()
def auth_headers(client: TestClient) -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        json={"username": "admin", "password": "admin12345"},
    )
    token = response.json()["data"]["tokens"]["access_token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture()
def sample_docx_bytes() -> bytes:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Header original"
    document.add_paragraph("Contenido base de la póliza")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Celda 1"
    table.cell(0, 1).text = "Celda 2"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture()
def settings_obj():
    from app.core.config import get_settings

    return get_settings()
