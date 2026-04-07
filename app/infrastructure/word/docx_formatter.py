from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.core.exceptions import DocumentProcessingError
from app.domain.products.entities import Product

TOKEN_PATTERN = re.compile(r"{{\s*(\w+)\s*}}")
PAGE_SIZE_MAP = {
    "A4": (21.0, 29.7),
    "LETTER": (21.59, 27.94),
}
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxFormatter:
    def format_document(
        self,
        *,
        source_path: Path,
        destination_path: Path,
        product: Product,
        policy_number: str,
    ) -> None:
        try:
            document = Document(source_path)
        except Exception as exc:
            raise DocumentProcessingError(
                "No fue posible abrir el archivo DOCX original.",
                code="docx_open_failed",
            ) from exc

        if product.active_format_rule is None:
            raise DocumentProcessingError(
                "El producto no tiene una regla activa para formateo.",
                code="missing_format_rule",
            )

        config = product.active_format_rule.configuration
        rendered_title = self._render_template(
            product.title_template,
            {"numero_poliza": policy_number, "producto": product.name},
        )
        rendered_header = self._render_template(
            product.header_template,
            {"numero_poliza": policy_number, "producto": product.name},
        )

        self._apply_page_setup(document, config.get("page_setup", {}))
        self._apply_general_text_format(document, config.get("general_text", {}))
        self._apply_headers(document, rendered_header)
        self._insert_title(document, rendered_title, config.get("title_text", {}))

        try:
            destination_path.parent.mkdir(parents=True, exist_ok=True)
            document.save(destination_path)
        except Exception as exc:
            raise DocumentProcessingError(
                "No fue posible guardar el documento formateado.",
                code="docx_save_failed",
            ) from exc

    def _apply_page_setup(self, document: DocumentType, page_setup: dict) -> None:
        paper_size = str(page_setup.get("paper_size", "A4")).upper()
        margins = page_setup.get("margins", {})
        width_height = PAGE_SIZE_MAP.get(paper_size)
        for section in document.sections:
            if width_height:
                section.page_width = Cm(width_height[0])
                section.page_height = Cm(width_height[1])
            if "top_cm" in margins:
                section.top_margin = Cm(float(margins["top_cm"]))
            if "bottom_cm" in margins:
                section.bottom_margin = Cm(float(margins["bottom_cm"]))
            if "left_cm" in margins:
                section.left_margin = Cm(float(margins["left_cm"]))
            if "right_cm" in margins:
                section.right_margin = Cm(float(margins["right_cm"]))

    def _apply_general_text_format(self, document: DocumentType, general_text: dict) -> None:
        normal_style = document.styles["Normal"]
        color = self._parse_color(general_text.get("color_hex"))

        if general_text.get("font_family"):
            normal_style.font.name = str(general_text["font_family"])
        if general_text.get("font_size_pt"):
            normal_style.font.size = Pt(float(general_text["font_size_pt"]))
        if color:
            normal_style.font.color.rgb = color
        if general_text.get("line_spacing"):
            normal_style.paragraph_format.line_spacing = float(general_text["line_spacing"])

        uppercase = bool(general_text.get("uppercase", False))
        for paragraph in self._iter_all_paragraphs(document):
            if general_text.get("line_spacing"):
                paragraph.paragraph_format.line_spacing = float(general_text["line_spacing"])
            for run in paragraph.runs:
                if general_text.get("font_family"):
                    run.font.name = str(general_text["font_family"])
                if general_text.get("font_size_pt"):
                    run.font.size = Pt(float(general_text["font_size_pt"]))
                if color:
                    run.font.color.rgb = color
                if uppercase and run.text:
                    run.text = run.text.upper()

    def _apply_headers(self, document: DocumentType, header_text: str) -> None:
        for section in document.sections:
            headers = [section.header]
            if getattr(section, "different_first_page_header_footer", False):
                headers.append(section.first_page_header)
            even_header = getattr(section, "even_page_header", None)
            if even_header is not None:
                headers.append(even_header)

            seen_ids = set()
            for header in headers:
                if id(header) in seen_ids:
                    continue
                seen_ids.add(id(header))
                try:
                    header.is_linked_to_previous = False
                except Exception:
                    pass
                self._replace_header_content(header, header_text)

    def _replace_header_content(self, header, header_text: str) -> None:
        paragraphs = list(header.paragraphs)
        if paragraphs:
            target = paragraphs[0]
            target.text = header_text
            target.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in paragraphs[1:]:
                paragraph._element.getparent().remove(paragraph._element)
        else:
            target = header.add_paragraph(header_text)
            target.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for table in list(header.tables):
            table._element.getparent().remove(table._element)

    def _insert_title(self, document: DocumentType, title: str, title_text: dict) -> None:
        first_non_empty = next((p for p in document.paragraphs if p.text.strip()), None)
        if first_non_empty and first_non_empty.text.strip() == title.strip():
            title_paragraph = first_non_empty
        else:
            title_paragraph = document.add_paragraph(title)
            body = document._body._element
            body.remove(title_paragraph._p)
            body.insert(0, title_paragraph._p)

        title_paragraph.style = document.styles["Title"]
        alignment = title_text.get("alignment")
        if alignment:
            title_paragraph.paragraph_format.alignment = ALIGNMENT_MAP.get(
                str(alignment).lower(),
                title_paragraph.paragraph_format.alignment,
            )

        for run in title_paragraph.runs:
            if title_text.get("font_family"):
                run.font.name = str(title_text["font_family"])
            if title_text.get("font_size_pt"):
                run.font.size = Pt(float(title_text["font_size_pt"]))
            if title_text.get("bold") is not None:
                run.bold = bool(title_text["bold"])
            if bool(title_text.get("uppercase", False)) and run.text:
                run.text = run.text.upper()

    def _iter_all_paragraphs(self, document: DocumentType) -> Iterable:
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

    def _render_template(self, template: str, context: dict[str, str]) -> str:
        def replace(match: re.Match[str]) -> str:
            key = match.group(1)
            return context.get(key, match.group(0))

        return TOKEN_PATTERN.sub(replace, template)

    def _parse_color(self, color_value: str | None):
        if not color_value:
            return None
        return RGBColor.from_string(str(color_value).replace("#", "").upper())
